#!/usr/bin/env python
# coding: utf-8

from pathlib import Path
from typing import Dict
import xml.dom.minidom
import zipfile
import os
# from docx.api import Document
import requests
# from spacy.language import Language
import extract.extractAdresse as extractAdresse
import tika.tika
from win32com import client as wc
import string
import re
import docx
import spacy
from HanTa import HanoverTagger as ht
from win32com import client as wc

tagger = ht.HanoverTagger('morphmodel_ger.pgz')


def docxPfad(pfad, datei):
    docxPfad = pfad + '\docx\\' + datei + 'x'
    currentPfad = pfad + '\\' + datei
    if datei.endswith('.doc'):
        if not os.path.exists(pfad + '\docx'):
            os.makedirs(pfad + '\docx')
        elif not os.path.exists(docxPfad):
            try:
                w = wc.Dispatch('Word.Application')
                doc = w.Documents.Open(currentPfad)
                doc.SaveAs(docxPfad, 16)  # Must have parameter 16, otherwise an error will occur.
                doc.Close()
                return docxPfad
            except:
                return ''
        elif os.path.exists(docxPfad):
            return docxPfad
        else:
            return ''
    elif datei.endswith('.docx'):
        return currentPfad
    else:
        return ''


def getTextContent(metadata: Dict, parser='tika', docxExists: bool = True) -> str:
    """
    :param dict metadata: The full metadata dictionary
    :param str parser: One of ["tika", "docx"]
    :param bool docxExists: Only needed if `methode` is docx
    :return str content: The file contents as a single string
    """
    if isinstance(metadata, dict):
        pfadOriginal = next(iter(metadata))
        datei = next(iter(metadata[pfadOriginal]))
        # pfad = metadata[pfadOriginal][datei]['pfadAktuell']  # Not req w/ relative folder struct
        pfad = pfadOriginal
    # if isinstance(metadata, list):
    #     # List mit Datei und Ordner, in dem es sich befindet
    #     pfad = metadata[1]
    #     datei = metadata[0]

    datei = Path(datei)
    pfad = Path(pfad)
    content = ''
    if parser == 'tika':
        try:
            filepath = pfad / datei
            content = extract_text(filepath, "http://localhost:9998")
            meta = extract_meta(filepath, "http://localhost:9998")
            print(meta)
        except:  # TODO Exception clause too broad. Raise for now and improve as errors come up
            content = ''
            raise

    # elif parser == 'tika_local':
    #     from tika import parser
    #     raw = parser.from_file(pfad / datei)
    #     content = raw['content']

    elif parser == 'docx':
        try:
            if not docxExists:
                if docxPfad(pfad, datei) != '':
                    doc = docx.Document(docxPfad(pfad, datei))
                    fullText = []
                    for para in doc.paragraphs:
                        fullText.append(para.text)
                    content = '\n'.join(fullText)
            else:
                if datei[-4:] == 'docx':
                    # Documents have already been converted to docx
                    doc = docx.Document(pfad + '\\' + datei)
                    fullText = []
                    for para in doc.paragraphs:
                        fullText.append(para.text)
                    content = '\n'.join(fullText)
                else:
                    content = ''
        except:  # TODO Exception clause too broad. Raise for now and improve as errors come up
            content = ''
            raise

    if content is None:
        content = ''

    return content


def getPageNumber(pfad, datei, methode, docxVorhanden=True):
    # TODO The if <methode == 'tika'> block was commented out by Christian. Why?
    # if methode == 'tika':
    #     from tika import parser
    #     file = pfad + '\\' + datei
    #     raw = parser.from_file(file)
    #     try:
    #         pages = raw['metadata']['xmpTPg:NPages']
    #         if type(pages) == str:
    #             pages = int(pages)
    #         else:
    #             pages = int(max(pages))
    #     except:
    #         pages = '1'
    if methode == 'docx':
        if not docxVorhanden:
            document = zipfile.ZipFile(docxPfad(pfad, datei))
        else:
            document = zipfile.ZipFile(pfad + '\\' + datei)
        dxml = document.read('docProps/app.xml')
        uglyXml = xml.dom.minidom.parseString(dxml)
        pages = uglyXml.getElementsByTagName('Pages')[0].childNodes[0].nodeValue
    else:
        pages = '1'
    return int(pages)


def getLemma(wordList, origWord):
    lemmaList = []
    for (word, lemma, pos) in tagger.tag_sent(wordList):
        possibleTags = [i[0] for i in tagger.tag_word(lemma, casesensitive=False, cutoff=5)]

        if possibleTags[0] in ['NE', 'ADJA', 'ADJD']:
            continue

        elif possibleTags[0] in ['ART']:
            if len(wordList) > 1:
                # charSplit was used but went wrong; get original word
                return [origWord]
            else:
                continue

        elif "denkmal" in lemma.lower():
            continue

        else:
            lemmaList.append(lemma)

    return lemmaList


def getLemmaRemvStopPunct(nlpdoc, stop_words):
    lemma_list = []
    for token in nlpdoc:
        tkLemma = token.lemma_
        try:
            splitWord = char_split.split_compound(tkLemma)
            if splitWord[0][0] > 0.9:  # oder vielleicht höher? 0.9?

                wort1 = splitWord[0][1]
                wort2 = splitWord[0][2]

                wortList = [wort1, wort2]

                lemma = getLemma(wortList, tkLemma)
                lemma_list.extend(lemma)

            else:
                lemma = getLemma([tkLemma], tkLemma)
                lemma_list.extend(lemma)

        except:
            lemma = getLemma([tkLemma], tkLemma)
            lemma_list.extend(lemma)

    """
    lemma_list1 = []
    for token in doc:
        lemma_list1.append(token.lemma_)
    text_tokens = lemma_list1
    print(text_tokens)
    print(lemma_list)
    """

    # Stopwords entfernen
    # tokens_without_sw = [word for word in text_tokens if not word in all_stopwords]
    tokens_without_sw = [word for word in lemma_list if not word.lower() in stop_words]

    # Remove punctuation
    tokens_without_punct = [''.join(c for c in s if c not in string.punctuation)
                            for s in tokens_without_sw]
    # print(tokens_without_sw)

    return tokens_without_punct


# Test pre-processing
def preprocessText(text: str, adresseMode=False, locSearch=False):
    nlpsp = spacy.load('de_core_news_lg')

    # Alternative nltk: https://towardsdatascience.com/text-normalization-with-spacy-and-nltk-1302ff430119

    # Überflüssige Leerzeichen im Text entfernen
    textCleanSpace = re.sub(' +', ' ',
                            text.replace('\n', ' ').replace('\t', ' ').replace('\r', ' ').rstrip())

    adressen = []
    if adresseMode:
        adressen = extractAdresse.getAddress(text)

    doc = nlpsp(textCleanSpace)

    all_stopwords = nlpsp.Defaults.stop_words

    lemmas_without_sw = getLemmaRemvStopPunct(doc, all_stopwords)

    # Remove empty strings
    tokensNotEmpty = []
    for word in lemmas_without_sw:
        if len(word) >= 3:
            tokensNotEmpty.append(word.lower())
    # tokens_without_sw = list(filter(None, tokens_without_sw))

    locations = []
    if locSearch:
        locations = list(set([ent.text for ent in doc.ents if ent.label_ in ['LOC']]))

    return tokensNotEmpty, adressen, locations


# TODO Why was this function changed from _extract_text() below?
def extract_text(file_path, tika_url):
    # TODO The opened file object potentially does not get cleaned properly
    response = requests.put(tika_url + "/tika", data=open(file_path, 'rb'))
    result = response.text
    return result


# TODO Why was this function changed from extract_meta() below?
def extract_meta(file_path, tika_url):
    file_path = Path(file_path)
    file_name = file_path.name
    # TODO The opened file object potentially does not get cleaned properly
    response = requests.put(tika_url + "/meta", data=open(file_path, 'rb'),
                            headers={"Accept": "application/json"})
    result = response.json()
    result['file_name'] = file_name
    return result

# async def _extract_text(file_path, tika_url):
#     async with aiohttp.ClientSession() as session:
#         async with session.put(url=tika_url, data=open(file_path, 'rb')) as response:
#             text_data = await response.text()
#             return text_data


# async def extract_meta(file_path, tika_url):
#     async with aiohttp.ClientSession() as session:
#         async with session.put(url=tika_url, data=open(file_path, 'rb'),
#                                headers={'Accept': 'application/json'}) as response:
#             file_name = str.split(file_path, '\\')[-1]
#             try:
#                 data = await response.text()
#                 if data:
#                     meta_data = json.loads(data)
#                     meta_data['file_name'] = file_name
#                 else:
#                     meta_data = ''
#             except Exception as ex:
#                 print("repsonse_failed")
#             return meta_data


# TODO Why was this function removed getTextContent() in favor of extract_text() & extract_meta()?
# def process_data(filepath: Path,
#                  tika_url: str = "http://localhost:9998/tika",
#                  content_type: str = 'text') -> str:
#     """
#
#     :param Path filepath: The full filepath to be parsed
#     :param str tika_url: The URL of the tika instance (running in docker)
#     :param str content_type: One of ["text", Any]; Parameter has no effect
#     :return str: The file contents
#     """
#     loop = asyncio.get_event_loop()
#     if content_type == 'text':
#         processed, unprocessed = loop.run_until_complete(
#             asyncio.wait([extract_text(j, tika_url) for j in [filepath]],
#                          timeout=60))
#     else:
#         processed, unprocessed = loop.run_until_complete(
#             asyncio.wait([extract_meta(j, tika_url) for j in [filepath]],
#                          timeout=60))
#     json_data = [f.result() for f in processed]
#     return json_data[0]
